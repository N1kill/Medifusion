import json
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# ---------------- PDF Report Generator ----------------


def generate_pdf_report():
    try:
        with open("chat_history.json", "r", encoding="utf-8") as f:
            all_history = json.load(f)

# keep only last 20 messages
        history = all_history[-20:]

    except:
        history = []

    filename = "report.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    y = 750

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 800, "Gemini AI Report")
    c.setFont("Helvetica", 12)
    c.drawString(50, 780, f"Generated on: {datetime.datetime.now()}")
    c.drawString(50, 765, "-" * 80)

    c.setFont("Helvetica", 11)

    for item in history:
        c.drawString(50, y, f"Prompt: {item['prompt']}")
        y -= 15
        c.drawString(50, y, f"Response: {item['response']}")
        y -= 30

        if y < 70:          # Start new page
            c.showPage()
            y = 750

    c.save()
    return filename

# ---------------- DOCX Report Generator ----------------

def generate_docx_report():
    try:
        with open("chat_history.json", "r") as f:
            history = json.load(f)
    except:
        history = []

    filename = "report.docx"
    doc = Document()

    doc.add_heading("Gemini AI Report", level=0)
    doc.add_paragraph(f"Generated on: {datetime.datetime.now()}")
    doc.add_paragraph("-" * 50)

    for item in history:
        doc.add_heading("Prompt:", level=1)
        doc.add_paragraph(item["prompt"])

        doc.add_heading("Response:", level=1)
        doc.add_paragraph(item["response"])

        doc.add_paragraph("")

    doc.save(filename)
    return filename

# ---------------- Logs ----------------

def append_log(prompt, response):
    with open("logs.txt", "a", encoding="utf-8") as log:
        log.write(f"\n[{datetime.datetime.now()}]\n")
        log.write(f"PROMPT: {prompt}\n")
        log.write(f"RESPONSE: {response}\n")
